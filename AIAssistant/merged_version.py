import http.client
import json
import re
from docx import Document
import OutputDocx
import PyPDF2

def read_pdf(file_path):
   with open(file_path, 'rb') as file:
       reader = PyPDF2.PdfReader(file)
       text = ''
       for page in reader.pages:
           text += page.extract_text()
   return text

def get_response(system_str, user_str, selected_modal='gpt-4o', max_tokens=1688):
   conn = http.client.HTTPSConnection("api.gpt.ge")

   payload = json.dumps({
      "model": selected_modal,
      "messages": [
         {
            "role": "system",
            "content": system_str
         },
         {
            "role": "user",
            "content": user_str
         }
      ],
      "max_tokens": max_tokens,
      "temperature": 1.0,
      "stream": False
   }) # temperature采样温度介于0和2之间。较高的值将使输出更加随机，而较低的值将使输出更加集中和确定。

   headers = {
      'Authorization': 'Bearer sk-5Q1tzButcM63EQ162bBb1f392eE54a85960d95B6C4Cb25B6',
      'Content-Type': 'application/json'
   }
   conn.request("POST", "/v1/chat/completions", payload, headers)
   res = conn.getresponse()
   data = res.read()
   return data

def find_keyword(text, keyword):
   pattern = re.compile(keyword)
   match = pattern.search(text)
   if match:
      return match.start()
   else:
      return f"Keyword '{keyword}' not found"

def read_docx(file_path):
   doc = Document(file_path)
   docx_content = ""
   for paragraph in doc.paragraphs:
      docx_content += paragraph.text + "\n"
   return docx_content

def read_docx2(file_path):
   try:
      doc = Document(file_path)
      if doc is None:
         return None
      docx_content = ""
      for paragraph in doc.paragraphs:
         docx_content += paragraph.text + "\n"
      return docx_content
   except:
      return None

def temporary_save(data):
   file1 = open("temporary.txt", "w", encoding='utf-8')
   file1.write(data.decode("utf-8"))
   file1.close()

def clean_content(text, pattern):
   matches = re.findall(pattern, text)
   cleaned_text = re.sub(pattern, '', text)
   return cleaned_text

segment_keywords = ["意义","现状","方向","文献"]

if __name__ == "__main__":
   template_str = read_docx('shenzhen_template.docx')
   start_i = find_keyword(template_str, segment_keywords[0])
   end_i = find_keyword(template_str, segment_keywords[1])
   template_str = template_str[start_i:end_i]
   template_str = clean_content(template_str, '\n')

   material_str = read_pdf('material1.pdf')
   start_i = find_keyword(material_str, segment_keywords[0])
   end_i = find_keyword(material_str, segment_keywords[1])
   material_str = material_str[start_i:end_i]
   material_str = clean_content(material_str, r'\b\d+草稿')
   material_str = clean_content(material_str, r'\d+/ - \d+ - ')
   material_str = clean_content(material_str, '\n')

   ''' SPECIFY HOW THE AI ASSISTANT SHOULD BEHAVE '''
   system_str = "Behave like a professor who can write grant proposal in Chinese"

   ''' SPECIFY WANT YOU WANT THE AI ASSISTANT TO SAY '''
   user_str = "以专业角度"
   user_str += "意义" \
               "每段"
   user_str += "\n模板如下：\n" + template_str
   user_str += "\n资料如下：\n" + material_str

   print(user_str)

   data = get_response(system_str, user_str, 'gpt-4o', 16384)
   print(data)
   temporary_save(data)
   OutputDocx.saveAsDocx("my")


   # for i in range(len(segment_keywords)-1):
   # material_str = read_pdf('D:/Code/SmushLBD/material1.pdf')
   # start_i = find_keyword(material_str, segment_keywords[i])
   # end_i = find_keyword(material_str, segment_keywords[i+1])
   # material_str = material_str[start_i:end_i]
   # material_str = clean_content(material_str, r'\b\d+草稿')
   # material_str = clean_content(material_str, r'\d+/ - \d+ - ')
   # material_str = clean_content(material_str, '\n')

   # from docx.oxml.ns import qn
   # docxName = "Material1"
   # doc = Document()
   # doc.styles['Normal'].font.name = u'宋体'
   # doc.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')
   # paragraph = doc.add_paragraph(material_str)
   # doc.save("%s.docx" % segment_keywords[i])